"""
Text Extraction Class

This module defines the TextExtractor class, which is responsible for
extracting text from PDF and DOCX files.

Author: Mirza Bilal
"""

import fitz
from docx import Document
from khulasa_document import KhDocument


class TextExtractor:
    """
    TextExtractor class for extracting text from PDF and DOCX files.

    Attributes:
        None

    Methods:
        __init__(): Initializes the TextExtractor.
        parse_file(file_path): Parses the input file and extracts text.
        _parse_pdf(pdf_path): Parses a PDF file and extracts text.
        _parse_docx(docx_file): Parses a DOCX file and extracts text.
    """

    # def __init__(self):
    #     """
    #     Initializes the TextExtractor.
    #     """
    #     pass

    def parse_file(self, file_path):
        """
        Parse the input file and extract text.

        :param file_path: The path to the input file.
        :return: A KhDocument object containing the extracted text.
        """
        file_extension = file_path.split(
            ".")[-1].lower()  # Extract and convert to lowercase

        # Check the file extension
        if file_extension == "pdf":
            return self._parse_pdf(file_path)

        elif file_extension == "docx":
            print("This is a DOCX file.")
            return self._parse_docx(file_path)
        else:
            print("Unsupported file format.")
            return None

    def _parse_pdf(self, pdf_path):
        """
        Parse a PDF file and extract text.

        :param pdf_path: The path to the PDF file.
        :return: A KhDocument object containing the extracted text.
        """
        pdf = fitz.open(pdf_path)
        document = KhDocument()
        current_page = None  # Initialize to None

        for page_number, page in enumerate(pdf, start=1):
            # Get text blocks
            blocks = page.get_text("blocks")

            for block in blocks:
                if '<image:' not in block[4]:
                    if current_page is None or current_page.page_number != page_number:
                        current_page = document.add_page(page_number)
                    current_page.add_block(block[4])
        return document

    def _parse_docx(self, docx_file):
        """
        Parse a DOCX file and extract text.

        :param docx_file: The path to the DOCX file.
        :return: A KhDocument object containing the extracted text.
        """
        document = Document(docx_file)
        kh_document = KhDocument()
        current_page = None

        for paragraph in document.paragraphs:
            if current_page is None:
                # Create a new page at the beginning of the document
                page_number = len(kh_document.pages) + 1
                current_page = kh_document.add_page(page_number)

            # Check if the paragraph represents a page break
            if paragraph.style.name == 'PageBreak':
                # Create a new page
                page_number = len(kh_document.pages) + 1
                current_page = kh_document.add_page(page_number)
            else:
                # Add the paragraph text as a block to the current page
                current_page.add_block(paragraph.text)

        return kh_document
